#!/usr/bin/env python3
"""
GDD Writer Template Generator for Infinite Voyage

Generates professional Game Design Document templates (.docx) including a cover
page, full GDD structure skeleton, and a style reference document. These templates
serve as the starting point for compiling polished, stakeholder-ready design documents.

Usage:
    python generate_templates.py
    python generate_templates.py --output-dir ./output

Requirements:
    pip install python-docx
"""

import argparse
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
except ImportError:
    print(
        "Error: python-docx is not installed.\n"
        "Install it with:\n"
        "    pip install python-docx\n"
        "\n"
        "Or if using a virtual environment:\n"
        "    python -m pip install python-docx"
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# Shared constants
# ---------------------------------------------------------------------------

BRAND_PRIMARY = RGBColor(0, 102, 255)  # Blue
BRAND_DARK = RGBColor(20, 20, 40)
BRAND_ACCENT = RGBColor(255, 153, 0)  # Orange
BRAND_GRAY = RGBColor(100, 100, 100)
BRAND_RED = RGBColor(200, 50, 50)
BRAND_GREEN = RGBColor(50, 150, 50)

TODAY_FORMATTED = datetime.now().strftime("%B %d, %Y")
TODAY_ISO = datetime.now().strftime("%Y-%m-%d")


def set_document_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set uniform margins on all sections (in inches)."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_page_number_footer(doc):
    """Add 'Page X of Y' footer to the document."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_para.add_run("Page ")

    # PAGE field
    run1 = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run1._element.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    run1._element.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run1._element.append(fld_end)

    footer_para.add_run(" of ")

    # NUMPAGES field
    run2 = footer_para.add_run()
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    run2._element.append(fld_begin2)

    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = "NUMPAGES \\* MERGEFORMAT"
    run2._element.append(instr2)

    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run2._element.append(fld_end2)


def add_centered_text(doc, text, font_size, bold=False, color=None, space_after=6):
    """Add a centered paragraph with specific formatting."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return para


# ===========================================================================
# 1. GDD Cover Page
# ===========================================================================


def generate_cover_page(output_dir):
    """Generate gdd-cover-page.docx with branded cover, metadata, and confidentiality."""
    doc = Document()
    set_document_margins(doc)

    # Vertical spacing before title
    for _ in range(5):
        doc.add_paragraph()

    # Game title
    add_centered_text(doc, "INFINITE VOYAGE", 48, bold=True, color=BRAND_PRIMARY)

    # Subtitle
    add_centered_text(
        doc, "Game Design Document", 24, bold=False, color=BRAND_GRAY, space_after=24
    )

    # Decorative separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("_" * 50)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Metadata table
    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    metadata = [
        ("Version", "[X.X]"),
        ("Date", TODAY_FORMATTED),
        ("Studio", "[Studio Name]"),
        ("Project Code", "[PROJECT_CODE]"),
        ("Lead Designer", "[Lead Designer Name]"),
        ("Status", "[Draft / Review / Final]"),
        ("Classification", "CONFIDENTIAL"),
    ]

    for i, (label, value) in enumerate(metadata):
        row = table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        label_cell.text = label
        for paragraph in label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

        value_cell.text = value
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

    # Spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Confidentiality notice
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf.paragraph_format.space_before = Pt(36)

    run_conf = conf.add_run("CONFIDENTIAL - INTERNAL USE ONLY")
    run_conf.font.size = Pt(14)
    run_conf.font.bold = True
    run_conf.font.color.rgb = BRAND_RED

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice.add_run(
        "This document contains proprietary information belonging to [Studio Name].\n"
        "Unauthorized reproduction, distribution, or disclosure is strictly prohibited.\n"
        "All rights reserved."
    )
    notice_run.font.size = Pt(9)
    notice_run.font.color.rgb = BRAND_GRAY

    # Page break after cover
    doc.add_page_break()

    # Second page: brief instructions
    doc.add_heading("How to Use This Template", level=1)
    doc.add_paragraph(
        "Replace all placeholder text in [brackets] with your project-specific information. "
        "Delete this instruction page before distributing the final document."
    )
    instructions = [
        "Update the version number on each revision.",
        "Fill in studio name, project code, and lead designer.",
        "Set the classification level appropriate for your audience.",
        "Replace the confidentiality notice with your legal department's approved text.",
        "Add your studio logo above the title if desired.",
    ]
    for item in instructions:
        doc.add_paragraph(item, style="List Bullet")

    filepath = os.path.join(output_dir, "gdd-cover-page.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ===========================================================================
# 2. GDD Structure (Full Skeleton)
# ===========================================================================


def generate_gdd_structure(output_dir):
    """Generate gdd-structure.docx with all section headings and placeholder text."""
    doc = Document()
    set_document_margins(doc)
    add_page_number_footer(doc)

    # Header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Infinite Voyage - Game Design Document"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = BRAND_GRAY

    # Title
    add_centered_text(doc, "INFINITE VOYAGE", 36, bold=True, color=BRAND_PRIMARY)
    add_centered_text(doc, "Game Design Document", 18, color=BRAND_GRAY, space_after=12)
    add_centered_text(doc, f"Version [X.X] | {TODAY_ISO}", 11, color=BRAND_GRAY)

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fld1 = parse_xml(
        r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w"))
    )
    instr = parse_xml(
        r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(
            nsdecls("w")
        )
    )
    fld2 = parse_xml(
        r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w"))
    )
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)

    doc.add_paragraph(
        "(Update this field in Word: select all with Ctrl+A, then press F9)",
        style="Normal",
    )

    doc.add_page_break()

    # ---------- FRONT MATTER ----------
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "[Write a 1-2 page overview of the game for busy stakeholders. Include the core "
        "pitch, target audience, key differentiators, and success criteria.]"
    )

    doc.add_heading("Quick Facts", level=2)
    # Quick facts table
    facts_table = doc.add_table(rows=8, cols=2)
    facts_table.style = "Light Grid Accent 1"
    facts_data = [
        ("Genre", "[e.g., Narrative Space Exploration RPG]"),
        ("Platform", "[e.g., PC, PS5, Xbox Series X/S]"),
        ("Target Audience", "[e.g., Ages 16+, fans of story-driven sci-fi]"),
        ("Estimated Play Time", "[e.g., 40-60 hours main story]"),
        ("Team Size", "[e.g., 15 core, 5 contractors]"),
        ("Budget", "[e.g., $2.5M]"),
        ("Development Timeline", "[e.g., 18 months to launch]"),
        ("Target Launch", "[e.g., Q4 2026]"),
    ]
    for i, (label, value) in enumerate(facts_data):
        facts_table.rows[i].cells[0].text = label
        facts_table.rows[i].cells[1].text = value

    doc.add_page_break()

    # ---------- MAIN CONTENT ----------
    sections = [
        (
            "1. High Concept",
            [
                ("1.1 Logline", "[One-sentence pitch that captures the essence of the game.]"),
                ("1.2 Setting", "[Describe the world, time period, and universe rules.]"),
                ("1.3 Core Premise", "[What is the player doing and why? What is the central conflict?]"),
                ("1.4 Target Audience", "[Primary and secondary audiences. Demographics, psychographics, comparable game fans.]"),
                ("1.5 Core Pillars", "[List 3-5 design pillars that guide every decision.]"),
            ],
        ),
        (
            "2. Gameplay",
            [
                ("2.1 Core Loop", "[Describe the primary gameplay loop: what the player does minute-to-minute.]"),
                ("2.2 Core Mechanics", "[Detail each major mechanic: movement, interaction, resource management.]"),
                ("2.3 Progression System", "[How does the player grow? Levels, skills, gear, story progression.]"),
                ("2.4 Combat / Interaction Model", "[Combat system, dialogue system, or primary interaction model.]"),
                ("2.5 Victory Conditions", "[Win states, fail states, and endgame content.]"),
                ("2.6 Controls & Input", "[Controller layout, key bindings, accessibility options.]"),
            ],
        ),
        (
            "3. Story & Narrative",
            [
                ("3.1 Story Overview", "[High-level story summary: beginning, middle, end.]"),
                ("3.2 Main Characters", "[Protagonist, antagonist, key NPCs with brief descriptions.]"),
                ("3.3 Dialogue & Tone", "[Writing style, humor level, darkness level, examples.]"),
                ("3.4 Narrative Flow", "[How story unfolds: linear, branching, open-world discovery.]"),
                ("3.5 Quests & Missions", "[Quest structure, side content, repeatable content.]"),
            ],
        ),
        (
            "4. Art & Visuals",
            [
                ("4.1 Visual Style", "[Art direction, reference images, mood boards.]"),
                ("4.2 Character Design", "[Character art guidelines, silhouette language, color coding.]"),
                ("4.3 Environment Design", "[Biome types, architectural styles, lighting direction.]"),
                ("4.4 UI/UX Style Guide", "[HUD layout, menu design, font choices, color palette.]"),
                ("4.5 Animation Guidelines", "[Animation priorities, style (realistic vs stylized), key animations.]"),
            ],
        ),
        (
            "5. Audio & Music",
            [
                ("5.1 Sound Direction", "[Overall audio aesthetic, reference tracks, mood.]"),
                ("5.2 Music Themes", "[Main theme, combat music, exploration, emotional moments.]"),
                ("5.3 Voice Acting", "[Voiced characters, casting direction, languages.]"),
                ("5.4 SFX Reference", "[Key sound effects, environmental audio, UI sounds.]"),
            ],
        ),
        (
            "6. Systems & Mechanics",
            [
                ("6.1 Inventory System", "[Inventory size, weight, sorting, quick-access.]"),
                ("6.2 Ability / Skill System", "[Skill trees, ability unlocks, active vs passive.]"),
                ("6.3 Equipment & Gear", "[Slots, rarity, upgrade paths, crafting integration.]"),
                ("6.4 Economy Model", "[Currency types, sinks and faucets, trading, pricing.]"),
                ("6.5 Crafting System", "[Materials, recipes, discovery, progression.]"),
                ("6.6 AI & Behavior", "[Enemy AI patterns, NPC schedules, companion behavior.]"),
            ],
        ),
        (
            "7. Level Design",
            [
                ("7.1 World Structure", "[Open world, hub-and-spoke, linear levels, procedural.]"),
                ("7.2 Level Progression", "[Difficulty curve, gating, exploration incentives.]"),
                ("7.3 Key Encounters", "[Boss fights, set pieces, memorable moments.]"),
                ("7.4 Level Layouts", "[Layout diagrams, flow charts, critical path descriptions.]"),
                ("7.5 Environmental Storytelling", "[How levels convey narrative without dialogue.]"),
            ],
        ),
        (
            "8. Technical Specifications",
            [
                ("8.1 Engine & Tools", "[Game engine, middleware, custom tools.]"),
                ("8.2 Performance Targets", "[FPS targets, resolution, load times, memory budget.]"),
                ("8.3 Platform-Specific Notes", "[PC requirements, console optimization, mobile considerations.]"),
                ("8.4 Networking & Multiplayer", "[Online features, co-op, leaderboards, cloud saves.]"),
                ("8.5 Accessibility", "[Subtitles, colorblind modes, remapping, difficulty options.]"),
            ],
        ),
    ]

    for chapter_title, subsections in sections:
        doc.add_heading(chapter_title, level=1)
        doc.add_paragraph(
            f"[Overview paragraph for {chapter_title}. Summarize this chapter's scope and key decisions.]"
        )
        for sub_title, placeholder in subsections:
            doc.add_heading(sub_title, level=2)
            doc.add_paragraph(placeholder)
        doc.add_page_break()

    # ---------- BACK MATTER ----------
    doc.add_heading("Appendix A: Glossary", level=1)
    doc.add_paragraph("[Define all game-specific terms, acronyms, and jargon used in this document.]")

    glossary_table = doc.add_table(rows=4, cols=2)
    glossary_table.style = "Light Grid Accent 1"
    glossary_table.rows[0].cells[0].text = "Term"
    glossary_table.rows[0].cells[1].text = "Definition"
    glossary_table.rows[1].cells[0].text = "[Term 1]"
    glossary_table.rows[1].cells[1].text = "[Definition]"
    glossary_table.rows[2].cells[0].text = "[Term 2]"
    glossary_table.rows[2].cells[1].text = "[Definition]"
    glossary_table.rows[3].cells[0].text = "[Term 3]"
    glossary_table.rows[3].cells[1].text = "[Definition]"

    doc.add_page_break()

    doc.add_heading("Appendix B: Reference Tables", level=1)
    doc.add_paragraph(
        "[Include stat tables, item databases, ability matrices, and economy data. "
        "Reference the data-modeler spreadsheets for full data.]"
    )

    doc.add_page_break()

    doc.add_heading("Appendix C: Change Log & Version History", level=1)
    history_table = doc.add_table(rows=4, cols=4)
    history_table.style = "Light Grid Accent 1"
    history_headers = ["Version", "Date", "Changes", "Author"]
    for i, h in enumerate(history_headers):
        history_table.rows[0].cells[i].text = h
    for row_idx in range(1, 4):
        history_table.rows[row_idx].cells[0].text = f"[{row_idx}.0]"
        history_table.rows[row_idx].cells[1].text = "[YYYY-MM-DD]"
        history_table.rows[row_idx].cells[2].text = "[Description of changes]"
        history_table.rows[row_idx].cells[3].text = "[Author]"

    doc.add_page_break()

    doc.add_heading("Appendix D: Approvals & Sign-Off", level=1)
    doc.add_paragraph(
        "The undersigned have reviewed and approved this Game Design Document."
    )

    approval_table = doc.add_table(rows=4, cols=4)
    approval_table.style = "Light Grid Accent 1"
    approval_headers = ["Role", "Name", "Signature", "Date"]
    for i, h in enumerate(approval_headers):
        approval_table.rows[0].cells[i].text = h
    roles = ["Lead Designer", "Art Director", "Technical Director"]
    for idx, role in enumerate(roles, 1):
        approval_table.rows[idx].cells[0].text = role
        approval_table.rows[idx].cells[1].text = "[Name]"
        approval_table.rows[idx].cells[2].text = "________________"
        approval_table.rows[idx].cells[3].text = "[Date]"

    doc.add_page_break()

    doc.add_heading("Distribution & Legal Notice", level=1)
    doc.add_paragraph(
        "This document is the property of [Studio Name]. It is provided for internal "
        "review purposes only and may not be reproduced, distributed, or disclosed to "
        "any third party without written consent from [Studio Name] leadership."
    )
    doc.add_paragraph(f"Document generated: {TODAY_ISO}")

    filepath = os.path.join(output_dir, "gdd-structure.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ===========================================================================
# 3. GDD Styles Reference
# ===========================================================================


def generate_gdd_styles(output_dir):
    """Generate gdd-styles.docx showing all custom styles with examples."""
    doc = Document()
    set_document_margins(doc)
    add_page_number_footer(doc)

    # ----- Define custom styles -----

    styles = doc.styles

    # -- DesignNote style --
    design_note_style = styles.add_style("DesignNote", WD_STYLE_TYPE.PARAGRAPH)
    design_note_style.font.size = Pt(10)
    design_note_style.font.italic = True
    design_note_style.font.color.rgb = RGBColor(0, 102, 180)
    design_note_style.paragraph_format.left_indent = Inches(0.5)
    design_note_style.paragraph_format.space_before = Pt(6)
    design_note_style.paragraph_format.space_after = Pt(6)

    # -- BalanceData style --
    balance_data_style = styles.add_style("BalanceData", WD_STYLE_TYPE.PARAGRAPH)
    balance_data_style.font.size = Pt(9)
    balance_data_style.font.name = "Courier New"
    balance_data_style.font.color.rgb = RGBColor(50, 100, 50)
    balance_data_style.paragraph_format.space_before = Pt(3)
    balance_data_style.paragraph_format.space_after = Pt(3)

    # -- StatusFinal style --
    status_final_style = styles.add_style("StatusFinal", WD_STYLE_TYPE.PARAGRAPH)
    status_final_style.font.size = Pt(10)
    status_final_style.font.bold = True
    status_final_style.font.color.rgb = RGBColor(34, 139, 34)

    # -- StatusDraft style --
    status_draft_style = styles.add_style("StatusDraft", WD_STYLE_TYPE.PARAGRAPH)
    status_draft_style.font.size = Pt(10)
    status_draft_style.font.bold = True
    status_draft_style.font.color.rgb = RGBColor(200, 150, 0)

    # -- StatusConcept style --
    status_concept_style = styles.add_style("StatusConcept", WD_STYLE_TYPE.PARAGRAPH)
    status_concept_style.font.size = Pt(10)
    status_concept_style.font.bold = True
    status_concept_style.font.color.rgb = RGBColor(200, 50, 50)

    # -- Callout style --
    callout_style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout_style.font.size = Pt(11)
    callout_style.font.bold = True
    callout_style.font.color.rgb = BRAND_ACCENT
    callout_style.paragraph_format.left_indent = Inches(0.3)
    callout_style.paragraph_format.space_before = Pt(12)
    callout_style.paragraph_format.space_after = Pt(12)

    # -- MetaInfo style --
    meta_style = styles.add_style("MetaInfo", WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.size = Pt(9)
    meta_style.font.italic = True
    meta_style.font.color.rgb = RGBColor(150, 150, 150)

    # ----- Build reference document -----

    add_centered_text(doc, "GDD Style Reference", 36, bold=True, color=BRAND_PRIMARY)
    add_centered_text(
        doc,
        "Visual guide to all custom styles used in Infinite Voyage GDD documents",
        14,
        color=BRAND_GRAY,
    )
    add_centered_text(doc, f"Generated: {TODAY_ISO}", 10, color=BRAND_GRAY)

    doc.add_page_break()

    # ---- Heading Styles ----
    doc.add_heading("1. Heading Styles", level=1)
    doc.add_paragraph(
        "Headings provide visual hierarchy and generate the Table of Contents. "
        "Use Heading 1 for chapters, Heading 2 for sections, Heading 3 for subsections, "
        "and Heading 4 for detail items."
    )

    doc.add_heading("Heading 1 - Chapter Level", level=1)
    doc.add_paragraph(
        "Use for major document chapters (e.g., 'High Concept', 'Gameplay', 'Systems')."
    )

    doc.add_heading("Heading 2 - Section Level", level=2)
    doc.add_paragraph(
        "Use for sections within chapters (e.g., 'Core Loop', 'Progression System')."
    )

    doc.add_heading("Heading 3 - Subsection Level", level=3)
    doc.add_paragraph(
        "Use for detailed topics within sections (e.g., 'Melee Combat', 'Ranged Combat')."
    )

    doc.add_heading("Heading 4 - Detail Level", level=4)
    doc.add_paragraph(
        "Use for specific items or notes (e.g., 'Sword Attack Animation', 'Dodge Mechanic')."
    )

    doc.add_page_break()

    # ---- Body Text Styles ----
    doc.add_heading("2. Body Text Styles", level=1)

    doc.add_heading("Normal / Body Text", level=2)
    doc.add_paragraph(
        "This is the default body text style used for all standard paragraphs in the GDD. "
        "It should be easy to read at length. Use 11pt font with 1.15 line spacing. "
        "Paragraphs are separated by 6pt spacing."
    )

    doc.add_heading("List Bullet", level=2)
    doc.add_paragraph("First bullet point: Used for listing features, requirements, or options.", style="List Bullet")
    doc.add_paragraph("Second bullet point: Keep bullets concise -- one idea per line.", style="List Bullet")
    doc.add_paragraph("Third bullet point: Use sub-bullets (List Bullet 2) for nested items.", style="List Bullet")

    doc.add_heading("List Number", level=2)
    doc.add_paragraph("First numbered item: Used for sequential steps or prioritized lists.", style="List Number")
    doc.add_paragraph("Second numbered item: Numbers imply order or priority.", style="List Number")
    doc.add_paragraph("Third numbered item: Use when sequence matters.", style="List Number")

    doc.add_page_break()

    # ---- Custom Styles ----
    doc.add_heading("3. Custom GDD Styles", level=1)

    doc.add_heading("DesignNote", level=2)
    doc.add_paragraph(
        "Use the DesignNote style for designer commentary that explains reasoning "
        "behind a design decision. These notes are indented and italicized in blue."
    )
    doc.add_paragraph(
        "DESIGN NOTE: The mage's low HP is intentional to create a risk-reward dynamic. "
        "Players who master positioning gain the highest DPS, while those who play "
        "recklessly are punished. This encourages skill-based play.",
        style="DesignNote",
    )

    doc.add_heading("BalanceData", level=2)
    doc.add_paragraph(
        "Use the BalanceData style for inline numeric data, formulas, or balance "
        "references. Displayed in monospace green font."
    )
    doc.add_paragraph(
        "Total_Health = Base_HP + (HP_Per_Level * (Level - 1))", style="BalanceData"
    )
    doc.add_paragraph(
        "Warrior Lv10: 100 + (8 * 9) = 172 HP", style="BalanceData"
    )
    doc.add_paragraph(
        "Mage Lv10: 60 + (5 * 9) = 105 HP", style="BalanceData"
    )

    doc.add_heading("StatusFinal", level=2)
    doc.add_paragraph(
        "Use status styles to indicate section completeness at a glance."
    )
    doc.add_paragraph("[FINAL] This section has been reviewed and approved.", style="StatusFinal")

    doc.add_heading("StatusDraft", level=2)
    doc.add_paragraph("[DRAFT] This section is in progress and may change.", style="StatusDraft")

    doc.add_heading("StatusConcept", level=2)
    doc.add_paragraph(
        "[CONCEPT] This section is at the concept stage and not yet implemented.", style="StatusConcept"
    )

    doc.add_heading("Callout", level=2)
    doc.add_paragraph(
        "Use the Callout style for important warnings, key decisions, or highlighted information."
    )
    doc.add_paragraph(
        "IMPORTANT: All damage values in this document assume a level 1 baseline. "
        "Scaling formulas are documented in the data-modeler spreadsheets.",
        style="Callout",
    )

    doc.add_heading("MetaInfo", level=2)
    doc.add_paragraph(
        "Use MetaInfo for section metadata like author attribution and last-updated dates."
    )
    doc.add_paragraph(
        f"Author: Sarah Chen | Last Updated: {TODAY_ISO} | Status: Final",
        style="MetaInfo",
    )

    doc.add_page_break()

    # ---- Table Styles ----
    doc.add_heading("4. Table Styles", level=1)
    doc.add_paragraph(
        "Tables in the GDD use Word's built-in 'Light Grid Accent 1' style for "
        "consistency. Below is an example data table."
    )

    table = doc.add_table(rows=4, cols=4)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Character"
    table.rows[0].cells[1].text = "HP"
    table.rows[0].cells[2].text = "Damage"
    table.rows[0].cells[3].text = "Speed"
    sample_data = [
        ("Warrior", "120", "28", "10"),
        ("Mage", "65", "12", "11"),
        ("Rogue", "80", "22", "16"),
    ]
    for i, (name, hp, dmg, spd) in enumerate(sample_data, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = hp
        table.rows[i].cells[2].text = dmg
        table.rows[i].cells[3].text = spd

    doc.add_page_break()

    # ---- Usage Guidelines ----
    doc.add_heading("5. Usage Guidelines", level=1)

    guidelines = [
        "Always use Heading styles (not manual bold/size changes) for section headers.",
        "Use DesignNote for any commentary that explains WHY a decision was made.",
        "Use BalanceData when referencing specific numeric values or formulas.",
        "Apply status styles at the top of each section to indicate completeness.",
        "Use Callout sparingly -- only for critical information that must not be missed.",
        "Apply MetaInfo below section headings to track authorship and updates.",
        "Use tables for structured data; avoid inline data dumps in paragraphs.",
        "Keep paragraphs under 5 sentences. Use bullet points for lists of 3+ items.",
    ]
    for item in guidelines:
        doc.add_paragraph(item, style="List Bullet")

    filepath = os.path.join(output_dir, "gdd-styles.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ===========================================================================
# CLI entry point
# ===========================================================================


def main():
    parser = argparse.ArgumentParser(
        description="Generate Infinite Voyage GDD document templates (.docx).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python generate_templates.py\n"
            "  python generate_templates.py --output-dir ./output\n"
        ),
    )
    parser.add_argument(
        "--output-dir",
        default=os.path.dirname(os.path.abspath(__file__)),
        help="Directory where .docx files will be written (default: script directory).",
    )
    args = parser.parse_args()

    output_dir = os.path.abspath(args.output_dir)
    os.makedirs(output_dir, exist_ok=True)

    print(f"Generating GDD writer templates in: {output_dir}\n")

    generate_cover_page(output_dir)
    generate_gdd_structure(output_dir)
    generate_gdd_styles(output_dir)

    print("\nDone. All GDD writer templates generated successfully.")


if __name__ == "__main__":
    main()
