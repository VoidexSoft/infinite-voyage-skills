#!/usr/bin/env python3
"""
GDD Writer Template Generator for Infinite Voyage

Generates professional Game Design Document templates (.docx) including a cover
page, full GDD structure skeleton, and a style reference document. These templates
serve as the starting point for compiling polished, stakeholder-ready design documents.

Usage:
    python generate_templates.py
    python generate_templates.py --output-dir ./output

Requirements:
    pip install python-docx
"""

import argparse
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
except ImportError:
    print(
        "Error: python-docx is not installed.\n"
        "Install it with:\n"
        "    pip install python-docx\n"
        "\n"
        "Or if using a virtual environment:\n"
        "    python -m pip install python-docx"
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# Shared constants
# ---------------------------------------------------------------------------

BRAND_PRIMARY = RGBColor(0, 102, 255)  # Blue
BRAND_DARK = RGBColor(20, 20, 40)
BRAND_ACCENT = RGBColor(255, 153, 0)  # Orange
BRAND_GRAY = RGBColor(100, 100, 100)
BRAND_RED = RGBColor(200, 50, 50)
BRAND_GREEN = RGBColor(50, 150, 50)

TODAY_FORMATTED = datetime.now().strftime("%B %d, %Y")
TODAY_ISO = datetime.now().strftime("%Y-%m-%d")


def set_document_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """Set uniform margins on all sections (in inches)."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_page_number_footer(doc):
    """Add 'Page X of Y' footer to the document."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_para.add_run("Page ")

    run1 = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run1._element.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    run1._element.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run1._element.append(fld_end)

    footer_para.add_run(" of ")

    run2 = footer_para.add_run()
    fld_begin2 = OxmlElement("w:fldChar")
    fld_begin2.set(qn("w:fldCharType"), "begin")
    run2._element.append(fld_begin2)

    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = "NUMPAGES \\* MERGEFORMAT"
    run2._element.append(instr2)

    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run2._element.append(fld_end2)


def add_centered_text(doc, text, font_size, bold=False, color=None, space_after=6):
    """Add a centered paragraph with specific formatting."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return para


# ===========================================================================
# Template generators (cover page, structure, styles)
# ===========================================================================
# NOTE: Full implementations moved from templates/ to scripts/
# See the original generate_templates.py for complete code.
# This file is the canonical location for these generators.


def generate_cover_page(output_dir):
    """Generate gdd-cover-page.docx with branded cover, metadata, and confidentiality."""
    doc = Document()
    set_document_margins(doc)

    for _ in range(5):
        doc.add_paragraph()

    add_centered_text(doc, "INFINITE VOYAGE", 48, bold=True, color=BRAND_PRIMARY)
    add_centered_text(doc, "Game Design Document", 24, bold=False, color=BRAND_GRAY, space_after=24)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("_" * 50)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.size = Pt(12)

    doc.add_paragraph()

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    metadata = [
        ("Version", "[X.X]"),
        ("Date", TODAY_FORMATTED),
        ("Studio", "[Studio Name]"),
        ("Project Code", "[PROJECT_CODE]"),
        ("Lead Designer", "[Lead Designer Name]"),
        ("Status", "[Draft / Review / Final]"),
        ("Classification", "CONFIDENTIAL"),
    ]

    for i, (label, value) in enumerate(metadata):
        row = table.rows[i]
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        label_cell.text = label
        for paragraph in label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        value_cell.text = value
        for paragraph in value_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf.paragraph_format.space_before = Pt(36)
    run_conf = conf.add_run("CONFIDENTIAL - INTERNAL USE ONLY")
    run_conf.font.size = Pt(14)
    run_conf.font.bold = True
    run_conf.font.color.rgb = BRAND_RED

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice.add_run(
        "This document contains proprietary information belonging to [Studio Name].\n"
        "Unauthorized reproduction, distribution, or disclosure is strictly prohibited.\n"
        "All rights reserved."
    )
    notice_run.font.size = Pt(9)
    notice_run.font.color.rgb = BRAND_GRAY

    doc.add_page_break()

    doc.add_heading("How to Use This Template", level=1)
    doc.add_paragraph(
        "Replace all placeholder text in [brackets] with your project-specific information. "
        "Delete this instruction page before distributing the final document."
    )
    instructions = [
        "Update the version number on each revision.",
        "Fill in studio name, project code, and lead designer.",
        "Set the classification level appropriate for your audience.",
        "Replace the confidentiality notice with your legal department's approved text.",
        "Add your studio logo above the title if desired.",
    ]
    for item in instructions:
        doc.add_paragraph(item, style="List Bullet")

    filepath = os.path.join(output_dir, "gdd-cover-page.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


def generate_gdd_structure(output_dir):
    """Generate gdd-structure.docx with all section headings and placeholder text."""
    doc = Document()
    set_document_margins(doc)
    add_page_number_footer(doc)

    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Infinite Voyage - Game Design Document"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = BRAND_GRAY

    add_centered_text(doc, "INFINITE VOYAGE", 36, bold=True, color=BRAND_PRIMARY)
    add_centered_text(doc, "Game Design Document", 18, color=BRAND_GRAY, space_after=12)
    add_centered_text(doc, f"Version [X.X] | {TODAY_ISO}", 11, color=BRAND_GRAY)

    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fld1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w")))
    instr = parse_xml(
        r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(nsdecls("w"))
    )
    fld2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w")))
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)

    doc.add_paragraph(
        "(Update this field in Word: select all with Ctrl+A, then press F9)",
        style="Normal",
    )

    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "[Write a 1-2 page overview of the game for busy stakeholders.]"
    )

    sections = [
        ("1. High Concept", [
            ("1.1 Logline", "[One-sentence pitch.]"),
            ("1.2 Setting", "[Describe the world.]"),
            ("1.3 Core Pillars", "[List 3-5 design pillars.]"),
        ]),
        ("2. Gameplay", [
            ("2.1 Core Loop", "[Primary gameplay loop.]"),
            ("2.2 Core Mechanics", "[Major mechanics.]"),
            ("2.3 Progression System", "[How the player grows.]"),
        ]),
        ("3. Story & Narrative", [
            ("3.1 Story Overview", "[High-level story summary.]"),
            ("3.2 Main Characters", "[Protagonist, antagonist, key NPCs.]"),
            ("3.3 Quests & Missions", "[Quest structure.]"),
        ]),
        ("4. Systems & Economy", [
            ("4.1 Economy Model", "[Currency types, sinks and faucets.]"),
            ("4.2 Equipment & Gear", "[Slots, rarity, upgrades.]"),
            ("4.3 Crafting System", "[Materials, recipes.]"),
        ]),
        ("5. Level Design", [
            ("5.1 World Structure", "[World layout.]"),
            ("5.2 Key Encounters", "[Boss fights, set pieces.]"),
        ]),
        ("6. Art & Audio", [
            ("6.1 Visual Style", "[Art direction.]"),
            ("6.2 Sound Direction", "[Audio aesthetic.]"),
        ]),
        ("7. Technical", [
            ("7.1 Engine & Tools", "[Game engine.]"),
            ("7.2 Performance Targets", "[FPS, resolution.]"),
        ]),
    ]

    for chapter_title, subsections in sections:
        doc.add_heading(chapter_title, level=1)
        for sub_title, placeholder in subsections:
            doc.add_heading(sub_title, level=2)
            doc.add_paragraph(placeholder)
        doc.add_page_break()

    doc.add_heading("Appendix: Change Log", level=1)
    doc.add_paragraph("[Version history here.]")

    filepath = os.path.join(output_dir, "gdd-structure.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


def generate_gdd_styles(output_dir):
    """Generate gdd-styles.docx showing all custom styles with examples."""
    doc = Document()
    set_document_margins(doc)

    add_centered_text(doc, "GDD Style Reference", 36, bold=True, color=BRAND_PRIMARY)
    add_centered_text(doc, f"Generated: {TODAY_ISO}", 10, color=BRAND_GRAY)

    doc.add_heading("Heading Styles", level=1)
    doc.add_paragraph("Use Heading 1-4 for document hierarchy.")

    doc.add_heading("Body Text", level=1)
    doc.add_paragraph("Standard body text for all paragraphs.")
    doc.add_paragraph("Bullet point example.", style="List Bullet")

    filepath = os.path.join(output_dir, "gdd-styles.docx")
    doc.save(filepath)
    print(f"  Created: {filepath}")


# ===========================================================================
# CLI entry point
# ===========================================================================


def main():
    parser = argparse.ArgumentParser(
        description="Generate Infinite Voyage GDD document templates (.docx).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python generate_templates.py\n"
            "  python generate_templates.py --output-dir ./output\n"
        ),
    )
    parser.add_argument(
        "--output-dir",
        default=os.path.dirname(os.path.abspath(__file__)),
        help="Directory where .docx files will be written (default: script directory).",
    )
    args = parser.parse_args()

    output_dir = os.path.abspath(args.output_dir)
    os.makedirs(output_dir, exist_ok=True)

    print(f"Generating GDD writer templates in: {output_dir}\n")

    generate_cover_page(output_dir)
    generate_gdd_structure(output_dir)
    generate_gdd_styles(output_dir)

    print("\nDone. All GDD writer templates generated successfully.")


if __name__ == "__main__":
    main()
